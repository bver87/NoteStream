import logging
import os

from docx import Document

log = logging.getLogger("notestream.agenda")


def extract_agenda_text(agenda_text: str | None, agenda_file_path: str | None) -> str | None:
    """
    Normalise agenda input to plain text.
    Priority: pasted text > uploaded file > None.
    """

    # 1. Pasted text takes priority
    if agenda_text and agenda_text.strip():
        return agenda_text.strip()

    # 2. Uploaded file
    if agenda_file_path and os.path.exists(agenda_file_path):
        ext = os.path.splitext(agenda_file_path)[1].lower()

        if ext == ".txt":
            with open(agenda_file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read().strip()
                log.info("Agenda loaded from TXT (%d chars)", len(text))
                return text or None

        if ext == ".docx":
            return _extract_docx(agenda_file_path)

        log.warning("Unsupported agenda file type: %s", ext)

    return None


def _extract_docx(path: str) -> str | None:
    """Extract all text from a DOCX file including table cells."""
    doc   = Document(path)
    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables — common in Dutch agenda formats
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    result = "\n".join(parts) if parts else None
    log.info("Agenda loaded from DOCX (%d parts)", len(parts))
    return result